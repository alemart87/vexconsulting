"""Render de exportación: Markdown → HTML con identidad Voicenter → DOCX/PDF.

Documento final paginado: portada, índice (con número de página en PDF),
numeración «Página X de Y» al pie y sangría francesa APA en Referencias.

Corre dentro de un subproceso (export_worker): imports perezosos y funciones
síncronas. En Windows dev WeasyPrint puede no estar disponible (GTK); el
worker traduce el error a un mensaje claro — en Docker/Render funciona.
"""
from __future__ import annotations

import re
from pathlib import Path

BRAND_CSS = """
@page {
  size: A4;
  margin: 22mm 18mm 20mm 18mm;
  @bottom-right { content: "Página " counter(page) " de " counter(pages); font-size: 8pt; color: #5B6275; }
  @bottom-left { content: "VEX Consulting · Voicenter S.A."; font-size: 8pt; color: #5B6275; }
}
@page cover { @bottom-right { content: none; } @bottom-left { content: none; } }
body { font-family: 'DejaVu Sans', sans-serif; font-size: 10.5pt; color: #2A2F3A; line-height: 1.55; }
h1, h2, h3 { color: #0F1116; line-height: 1.15; page-break-after: avoid; }
h1 { font-size: 21pt; text-transform: uppercase; border-bottom: 3px solid #E6332A; padding-bottom: 6px; }
h2 { font-size: 15pt; text-transform: uppercase; margin-top: 22px; }
h3 { font-size: 12pt; margin-top: 16px; }
table { border-collapse: collapse; width: 100%; margin: 12px 0; font-size: 9pt; page-break-inside: avoid; }
th { background: #F6F7FB; border: 1px solid #E5E7EB; padding: 6px 8px; text-align: left; }
td { border: 1px solid #E5E7EB; padding: 6px 8px; }
blockquote { border-left: 4px solid #E6332A; margin: 12px 0; padding: 4px 14px; color: #5B6275; font-style: italic; }
code { background: #F6F7FB; padding: 1px 4px; border-radius: 3px; font-size: 9pt; color: #662483; }
pre { background: #0F1116; color: #fff; padding: 10px; border-radius: 6px; font-size: 8.5pt; overflow-x: auto; }
img { max-width: 100%; }
a { color: #00B2BF; }
.cover { page: cover; page-break-after: always; text-align: left; padding-top: 60mm; }
.cover .band { height: 10mm; background: #E6332A; margin: 0 -18mm; }
.cover h1 { border: none; font-size: 28pt; margin-top: 26mm; }
.cover .meta { color: #5B6275; font-size: 10pt; margin-top: 8mm; }
.toc { page-break-after: always; }
.toc h2 { border-bottom: 2px solid #E6332A; padding-bottom: 4px; }
.toc p { margin: 4px 0; }
.toc a { color: #2A2F3A; text-decoration: none; }
.toc a::after { content: "  ·  p. " target-counter(attr(href), page); color: #5B6275; font-size: 9pt; }
.toc .toc-2 { padding-left: 6mm; }
.toc .toc-3 { padding-left: 12mm; font-size: 9.5pt; }
.refs p { padding-left: 12mm; text-indent: -12mm; margin: 7px 0; }
"""

_HEADING_RE = re.compile(r"<h([123])>(.*?)</h\1>", re.DOTALL)
_TAG_RE = re.compile(r"<[^>]+>")


def _anchor_headings(body: str) -> tuple[str, list[tuple[int, str, str]]]:
    """Agrega id a cada h1-h3 del cuerpo y devuelve (html, [(nivel, id, texto)])."""
    entries: list[tuple[int, str, str]] = []
    counter = {"n": 0}

    def repl(match: re.Match) -> str:
        counter["n"] += 1
        hid = f"sec-{counter['n']}"
        level, inner = int(match.group(1)), match.group(2)
        entries.append((level, hid, _TAG_RE.sub("", inner).strip()))
        return f'<h{level} id="{hid}">{inner}</h{level}>'

    return _HEADING_RE.sub(repl, body), entries


def _build_toc(entries: list[tuple[int, str, str]]) -> str:
    if len(entries) < 2:
        return ""
    items = "\n".join(
        f'<p class="toc-{level}"><a href="#{hid}">{text}</a></p>'
        for level, hid, text in entries
    )
    return f'<nav class="toc"><h2>Índice</h2>\n{items}\n</nav>'


def _wrap_references(body: str) -> str:
    """Envuelve la sección Referencias en div.refs (sangría francesa APA)."""
    match = re.search(r"<h2[^>]*>\s*Referencias\s*</h2>", body, re.IGNORECASE)
    if not match:
        return body
    head, rest = body[: match.end()], body[match.end():]
    nxt = re.search(r"<h[12][ >]", rest)
    if nxt:
        return head + f'<div class="refs">{rest[: nxt.start()]}</div>' + rest[nxt.start():]
    return head + f'<div class="refs">{rest}</div>'


def md_to_html(content_md: str, title: str, author_note: str, upload_root: str,
               project_id: str) -> str:
    from markdown_it import MarkdownIt

    md = MarkdownIt("commonmark", {"html": False, "linkify": True}).enable("table").enable("strikethrough")
    body = md.render(content_md or "")
    body = _resolve_images(body, upload_root, project_id)
    body, headings = _anchor_headings(body)
    body = _wrap_references(body)
    toc = _build_toc(headings)
    return f"""<!DOCTYPE html>
<html lang="es"><head><meta charset="utf-8"><title>{_esc(title)}</title>
<style>{BRAND_CSS}</style></head>
<body>
<div class="cover">
  <div class="band"></div>
  <h1>{_esc(title)}</h1>
  <p class="meta">{_esc(author_note)}</p>
  <p class="meta">VEX Consulting · Plataforma de investigación de mercado · Voicenter S.A.</p>
</div>
{toc}
{body}
</body></html>"""


def _esc(text: str) -> str:
    return (text or "").replace("&", "&amp;").replace("<", "&lt;").replace(">", "&gt;")


def _resolve_images(html: str, upload_root: str, project_id: str) -> str:
    """Convierte las URLs internas /api/v1/projects/{id}/images/{name} a rutas
    file:// locales para que el exportador las incruste."""
    pattern = re.compile(r'src="/api/v1/projects/([^/"]+)/images/([^"]+)"')

    def repl(match: re.Match) -> str:
        pid, name = match.group(1), match.group(2)
        if pid != project_id or "/" in name or ".." in name:
            return 'src=""'
        local = Path(upload_root) / pid / "images" / name
        return f'src="{local.as_uri()}"' if local.exists() else 'src=""'

    return pattern.sub(repl, html)


# ---------------------------------------------------------------------------
# DOCX: pandoc + post-proceso python-docx (portada, índice real y pie paginado)
# ---------------------------------------------------------------------------

def _docx_simple_field(paragraph, instruction: str, placeholder: str = "1") -> None:
    """Inserta un campo de Word (PAGE, NUMPAGES) al final del párrafo."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    fld = OxmlElement("w:fldSimple")
    fld.set(qn("w:instr"), instruction)
    run = OxmlElement("w:r")
    text = OxmlElement("w:t")
    text.text = placeholder
    run.append(text)
    fld.append(run)
    paragraph._p.append(fld)


def _docx_toc_field(paragraph) -> None:
    """Inserta un campo TOC REAL de Word (números de página correctos).

    Se marca dirty y, junto con updateFields en settings.xml, Word lo
    recalcula al abrir el documento — a diferencia del índice estático de
    pandoc, que mostraba «1» en todas las entradas."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    begin = OxmlElement("w:fldChar")
    begin.set(qn("w:fldCharType"), "begin")
    begin.set(qn("w:dirty"), "true")
    r = paragraph.add_run()
    r._r.append(begin)

    instr = OxmlElement("w:instrText")
    instr.set(qn("xml:space"), "preserve")
    instr.text = ' TOC \\o "1-3" \\h \\z \\u '
    r = paragraph.add_run()
    r._r.append(instr)

    sep = OxmlElement("w:fldChar")
    sep.set(qn("w:fldCharType"), "separate")
    r = paragraph.add_run()
    r._r.append(sep)

    paragraph.add_run(
        "El índice se genera al abrir el documento en Word "
        "(si no aparece: clic derecho → Actualizar campos)."
    )

    end = OxmlElement("w:fldChar")
    end.set(qn("w:fldCharType"), "end")
    r = paragraph.add_run()
    r._r.append(end)


def _docx_enable_update_fields(document) -> None:
    """settings.xml: Word actualiza los campos (índice, numeración) al abrir."""
    from docx.oxml import OxmlElement
    from docx.oxml.ns import qn

    settings = document.settings.element
    if settings.find(qn("w:updateFields")) is None:
        upd = OxmlElement("w:updateFields")
        upd.set(qn("w:val"), "true")
        settings.append(upd)


def _decorate_docx(path: str, title: str, author_note: str) -> None:
    """Agrega portada, índice real (campo TOC) y pie «Página X de Y»."""
    import docx
    from docx.enum.text import WD_ALIGN_PARAGRAPH, WD_BREAK, WD_TAB_ALIGNMENT
    from docx.shared import Cm, Pt, RGBColor

    d = docx.Document(path)
    ink = RGBColor(0x0F, 0x11, 0x16)
    slate = RGBColor(0x5B, 0x62, 0x75)
    red = RGBColor(0xE6, 0x33, 0x2A)

    # Portada + índice: se insertan ANTES del primer párrafo del cuerpo.
    if d.paragraphs:
        first = d.paragraphs[0]

        def before(text: str, size: int, bold: bool, color, space_before: int = 0):
            p = first.insert_paragraph_before(text)
            p.paragraph_format.space_before = Pt(space_before)
            run = p.runs[0] if p.runs else p.add_run("")
            run.font.size = Pt(size)
            run.font.bold = bold
            run.font.color.rgb = color
            return p

        before("VEX CONSULTING", 14, True, red, space_before=80)
        before(title, 28, True, ink, space_before=24)
        before(author_note, 10, False, slate, space_before=18)
        before("VEX Consulting · Plataforma de investigación de mercado · Voicenter S.A.",
               10, False, slate, space_before=4)
        brk = first.insert_paragraph_before("")
        brk.add_run().add_break(WD_BREAK.PAGE)

        # Índice real: título (fuera del TOC) + campo TOC + salto de página
        toc_title = before("Índice", 20, True, ink, space_before=8)
        toc_title.paragraph_format.space_before = Pt(0)
        toc_par = first.insert_paragraph_before("")
        _docx_toc_field(toc_par)
        brk2 = first.insert_paragraph_before("")
        brk2.add_run().add_break(WD_BREAK.PAGE)

    _docx_enable_update_fields(d)

    # Pie de página con numeración en todas las secciones.
    for section in d.sections:
        footer = section.footer
        footer.is_linked_to_previous = False
        p = footer.paragraphs[0] if footer.paragraphs else footer.add_paragraph()
        for run in list(p.runs):
            run.text = ""
        p.alignment = WD_ALIGN_PARAGRAPH.LEFT
        p.paragraph_format.tab_stops.add_tab_stop(Cm(17), WD_TAB_ALIGNMENT.RIGHT)
        p.add_run("VEX Consulting · Voicenter S.A.\tPágina ")
        _docx_simple_field(p, "PAGE")
        p.add_run(" de ")
        _docx_simple_field(p, "NUMPAGES")
        for run in p.runs:
            run.font.size = Pt(8)
            run.font.color.rgb = slate

    d.save(path)


def _drop_duplicate_title(md: str, title: str) -> str:
    """Quita el H1 inicial si repite el título del proyecto: la portada ya lo
    muestra y duplicarlo ensucia el índice y la primera página."""
    lines = (md or "").lstrip().splitlines()
    if lines and lines[0].startswith("# "):
        heading = lines[0][2:].strip().lower()
        if heading == (title or "").strip().lower():
            return "\n".join(lines[1:]).lstrip()
    return md or ""


def export_docx(content_md: str, title: str, author_note: str, output_path: str,
                upload_root: str, project_id: str) -> None:
    """Markdown → DOCX vía pandoc; portada, índice real (campo TOC) y paginado
    se agregan en post-proceso (python-docx)."""
    import pypandoc

    md = _drop_duplicate_title(content_md, title)

    def _img_path(m: re.Match) -> str:
        # SOLO imágenes del propio proyecto (misma validación que el PDF):
        # sin esto, una referencia a otro project_id incrustaría archivos ajenos.
        pid, name = m.group(1), m.group(2)
        if pid != project_id or "/" in name or "\\" in name or ".." in name:
            return "()"
        return f"({(Path(upload_root) / pid / 'images' / name).as_posix()})"

    md = re.sub(r"\(/api/v1/projects/([^/)]+)/images/([^)]+)\)", _img_path, md)
    pypandoc.convert_text(
        md,
        "docx",
        format="gfm",
        outputfile=output_path,
        extra_args=["--standalone", "-M", "lang=es"],
    )
    try:
        _decorate_docx(output_path, title, author_note)
    except Exception:
        # La portada/pie son decorativos: el DOCX de pandoc ya es válido.
        import logging

        logging.getLogger("vexconsulting").exception("No se pudo decorar el DOCX")


def export_pdf(content_md: str, title: str, author_note: str, output_path: str,
               upload_root: str, project_id: str) -> None:
    """Markdown → HTML de marca (portada + índice paginado) → PDF vía WeasyPrint."""
    from weasyprint import HTML

    content_md = _drop_duplicate_title(content_md, title)
    html = md_to_html(content_md, title, author_note, upload_root, project_id)
    HTML(string=html, base_url=upload_root).write_pdf(output_path)


# ---------------------------------------------------------------------------
# PAPER: publicación ligera (LinkedIn, clientes) — sin normas APA, sin índice.
# Portada con logo (Voicenter o personalizado), título grande y tarjeta de
# autor con foto abajo. Tipografía cómoda de lectura, marca sobria.
# ---------------------------------------------------------------------------

VOICENTER_LOGO = Path(__file__).resolve().parents[2] / "assets" / "logo-voicenter-color.png"

PAPER_CSS = """
@page {
  size: A4;
  margin: 20mm 20mm 20mm 20mm;
  @bottom-left { content: "__BRAND__"; font-size: 8pt; color: #9AA0AE; }
  @bottom-right { content: counter(page) " / " counter(pages); font-size: 8pt; color: #9AA0AE; }
}
@page cover { @bottom-left { content: none; } @bottom-right { content: none; } margin: 0; }
body { font-family: 'DejaVu Sans', sans-serif; font-size: 10.5pt; color: #2A2F3A; line-height: 1.7; }

/* ---- Portada ---- */
.paper-cover { page: cover; page-break-after: always; position: relative;
  height: 297mm; padding: 22mm 20mm 18mm 20mm; box-sizing: border-box; }
.paper-cover .logo { height: 16mm; max-width: 70mm; object-fit: contain; object-position: left; }
.paper-cover .accent { height: 1.6mm; width: 34mm; background: #E6332A; margin-top: 10mm; }
.paper-cover h1 { font-size: 30pt; line-height: 1.12; color: #0F1116; margin: 8mm 0 0 0;
  border: none; text-transform: none; font-weight: 800; }
.paper-cover .subtitle { font-size: 13pt; color: #5B6275; margin-top: 6mm; line-height: 1.5; }
.paper-cover .footer { position: absolute; bottom: 18mm; left: 20mm; right: 20mm; }
.paper-cover .author { display: flex; align-items: center; gap: 6mm;
  border-top: 0.4mm solid #E5E7EB; padding-top: 6mm; }
.paper-cover .author img.photo { width: 22mm; height: 22mm; border-radius: 50%;
  object-fit: cover; border: 1mm solid #fff; box-shadow: 0 1mm 3mm rgba(15,17,22,.18); }
.paper-cover .author .initials { width: 22mm; height: 22mm; border-radius: 50%;
  background: #662483; color: #fff; text-align: center; line-height: 22mm;
  font-size: 15pt; font-weight: 800; }
.paper-cover .author .name { font-size: 12.5pt; font-weight: 800; color: #0F1116; }
.paper-cover .author .role { font-size: 9.5pt; color: #5B6275; margin-top: 1mm; }
.paper-cover .date { font-size: 9pt; color: #9AA0AE; margin-top: 4mm; }

/* ---- Cuerpo ---- */
h1, h2, h3 { color: #0F1116; line-height: 1.2; page-break-after: avoid; font-weight: 800; }
h1 { font-size: 17pt; margin-top: 22px; border: none; text-transform: none; }
h2 { font-size: 14pt; margin-top: 20px; text-transform: none; }
h2::before { content: ""; display: inline-block; width: 5mm; height: 1.2mm;
  background: #E6332A; margin-right: 2.5mm; vertical-align: middle; }
h3 { font-size: 11.5pt; margin-top: 14px; }
p { margin: 8px 0; }
table { border-collapse: collapse; width: 100%; margin: 12px 0; font-size: 9pt; page-break-inside: avoid; }
th { background: #0F1116; color: #fff; border: 1px solid #0F1116; padding: 6px 8px; text-align: left; }
td { border: 1px solid #E5E7EB; padding: 6px 8px; }
tr:nth-child(even) td { background: #F9FAFC; }
blockquote { border-left: 3px solid #E6332A; margin: 14px 0; padding: 2px 14px;
  color: #414855; font-size: 11.5pt; }
img { max-width: 100%; }
a { color: #00B2BF; text-decoration: none; }
code { background: #F6F7FB; padding: 1px 4px; border-radius: 3px; font-size: 9pt; color: #662483; }

/* ---- Cierre con autor ---- */
.paper-end { margin-top: 14mm; border-top: 0.4mm solid #E5E7EB; padding-top: 6mm;
  display: flex; align-items: center; gap: 5mm; page-break-inside: avoid; }
.paper-end img.photo { width: 16mm; height: 16mm; border-radius: 50%; object-fit: cover; }
.paper-end .initials { width: 16mm; height: 16mm; border-radius: 50%; background: #662483;
  color: #fff; text-align: center; line-height: 16mm; font-size: 11pt; font-weight: 800; }
.paper-end .name { font-weight: 800; font-size: 10.5pt; color: #0F1116; }
.paper-end .role { font-size: 9pt; color: #5B6275; }
.paper-end img.logo { height: 9mm; max-width: 42mm; object-fit: contain; margin-left: auto; }
"""


def _author_avatar(photo_uri: str | None, autor: str, css_class: str) -> str:
    if photo_uri:
        return f'<img class="photo" src="{photo_uri}" alt="{_esc(autor)}"/>'
    initials = "".join(w[0] for w in (autor or "V").split()[:2]).upper()
    return f'<div class="initials {css_class}">{_esc(initials)}</div>'


def paper_html(content_md: str, opts: dict, upload_root: str, project_id: str) -> str:
    """HTML del paper: portada (logo + título + autor con foto abajo) + cuerpo."""
    from datetime import datetime, timezone

    from markdown_it import MarkdownIt

    titulo = str(opts.get("titulo") or "Paper")
    subtitulo = str(opts.get("subtitulo") or "")
    autor = str(opts.get("autor") or "")
    cargo = str(opts.get("cargo") or "")

    # Logo: Voicenter empaquetado, o el archivo subido por el consultor
    logo_uri = VOICENTER_LOGO.as_uri() if VOICENTER_LOGO.exists() else ""
    logo_name = str(opts.get("logo") or "voicenter")
    if logo_name != "voicenter" and "/" not in logo_name and ".." not in logo_name:
        custom = Path(upload_root) / project_id / "paper" / logo_name
        if custom.exists():
            logo_uri = custom.as_uri()

    photo_uri = None
    foto = str(opts.get("foto") or "")
    if foto and "/" not in foto and ".." not in foto:
        photo_path = Path(upload_root) / project_id / "paper" / foto
        if photo_path.exists():
            photo_uri = photo_path.as_uri()

    md = MarkdownIt("commonmark", {"html": False, "linkify": True}).enable("table").enable("strikethrough")
    body = _resolve_images(md.render(_drop_duplicate_title(content_md, titulo)), upload_root, project_id)

    meses = ("enero", "febrero", "marzo", "abril", "mayo", "junio", "julio",
             "agosto", "septiembre", "octubre", "noviembre", "diciembre")
    now = datetime.now(timezone.utc)
    fecha = f"{meses[now.month - 1].capitalize()} {now.year}"
    brand = autor or "Voicenter S.A."
    author_html = ""
    if autor:
        author_html = f"""
  <div class="author">
    {_author_avatar(photo_uri, autor, "")}
    <div>
      <div class="name">{_esc(autor)}</div>
      {f'<div class="role">{_esc(cargo)}</div>' if cargo else ''}
    </div>
  </div>"""

    end_html = ""
    if autor:
        end_html = f"""
<div class="paper-end">
  {_author_avatar(photo_uri, autor, "")}
  <div>
    <div class="name">{_esc(autor)}</div>
    {f'<div class="role">{_esc(cargo)}</div>' if cargo else ''}
  </div>
  {f'<img class="logo" src="{logo_uri}" alt="logo"/>' if logo_uri else ''}
</div>"""

    # La marca del pie va directo en el CSS (string-set de WeasyPrint es frágil)
    brand_css = PAPER_CSS.replace("__BRAND__", brand.replace('"', "'")[:80])

    return f"""<!DOCTYPE html>
<html lang="es"><head><meta charset="utf-8"><title>{_esc(titulo)}</title>
<style>{brand_css}</style></head>
<body>
<div class="paper-cover">
  {f'<img class="logo" src="{logo_uri}" alt="logo"/>' if logo_uri else ''}
  <div class="accent"></div>
  <h1>{_esc(titulo)}</h1>
  {f'<p class="subtitle">{_esc(subtitulo)}</p>' if subtitulo else ''}
  <div class="footer">{author_html}
    <p class="date">{_esc(fecha)}</p>
  </div>
</div>
{body}
{end_html}
</body></html>"""


def export_paper(content_md: str, opts: dict, output_path: str,
                 upload_root: str, project_id: str) -> None:
    """Paper ligero → PDF vía WeasyPrint."""
    from weasyprint import HTML

    html = paper_html(content_md, opts or {}, upload_root, project_id)
    HTML(string=html, base_url=upload_root).write_pdf(output_path)


def run_export(fmt: str, content_md: str, title: str, author_note: str,
               output_path: str, upload_root: str, project_id: str,
               options: dict | None = None) -> None:
    """Punto de entrada del subproceso."""
    if fmt == "docx":
        export_docx(content_md, title, author_note, output_path, upload_root, project_id)
    elif fmt == "pdf":
        export_pdf(content_md, title, author_note, output_path, upload_root, project_id)
    elif fmt == "paper":
        export_paper(content_md, {**(options or {}), "titulo": (options or {}).get("titulo") or title},
                     output_path, upload_root, project_id)
    else:
        raise ValueError(f"Formato no soportado: {fmt}")
