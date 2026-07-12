"""Extracción de texto de las fuentes (PDF, Word, Excel, links, texto plano).

Corre dentro de un subproceso aislado (jobs) — acá solo lógica pura y síncrona.
Cada extractor devuelve una lista de secciones: {text, meta{page|sheet|section}}.
"""
from __future__ import annotations

from pathlib import Path
from typing import Any

Section = dict[str, Any]  # {"text": str, "meta": {...}}

_MIN_CHARS_PER_PAGE = 100  # bajo esto, el PDF probablemente es escaneado


def extract_pdf(path: str) -> tuple[list[Section], int]:
    """Cascada: pypdf → pdfplumber → PyMuPDF → OCR con visión de GPT
    (para PDFs escaneados, si hay OPENAI_API_KEY y OCR habilitado)."""
    from pypdf import PdfReader

    reader = PdfReader(path)
    pages: list[Section] = []
    total_chars = 0
    for i, page in enumerate(reader.pages, start=1):
        text = (page.extract_text() or "").strip()
        total_chars += len(text)
        if text:
            pages.append({"text": text, "meta": {"page": i}})

    n_pages = len(reader.pages)

    def _is_poor() -> bool:
        chars = sum(len(p["text"]) for p in pages)
        return not n_pages or chars / max(n_pages, 1) < _MIN_CHARS_PER_PAGE

    if _is_poor():
        try:
            import pdfplumber

            alt: list[Section] = []
            with pdfplumber.open(path) as pdf:
                for i, page in enumerate(pdf.pages, start=1):
                    text = (page.extract_text() or "").strip()
                    if text:
                        alt.append({"text": text, "meta": {"page": i}})
            if sum(len(p["text"]) for p in alt) > sum(len(p["text"]) for p in pages):
                pages = alt
        except Exception:
            pass

    if _is_poor():
        try:
            alt = _pymupdf_text(path)
            if sum(len(p["text"]) for p in alt) > sum(len(p["text"]) for p in pages):
                pages = alt
        except Exception:
            pass

    if _is_poor():
        # PDF escaneado: OCR con visión de GPT (página → imagen → transcripción)
        from ...core.config import settings

        if settings.ocr_enabled and settings.openai_api_key:
            ocr_pages = _vision_ocr(path, max_pages=30)
            if ocr_pages:
                pages = ocr_pages

    return pages, n_pages


def _pymupdf_text(path: str) -> list[Section]:
    import fitz  # PyMuPDF

    doc = fitz.open(path)
    pages: list[Section] = []
    for i, page in enumerate(doc, start=1):
        text = (page.get_text() or "").strip()
        if text:
            pages.append({"text": text, "meta": {"page": i}})
    doc.close()
    return pages


_OCR_PROMPT = (
    "Transcribí TODO el texto visible de esta imagen, en orden de lectura y en su "
    "idioma original. Las tablas como tablas Markdown. Si hay gráficos, extraé sus "
    "datos y describilos brevemente. Respondé SOLO con el contenido; si no hay nada "
    "legible, respondé VACIO."
)


def _vision_available() -> bool:
    from ...core.config import settings

    return bool(settings.ocr_enabled and settings.openai_api_key)


def _vision_image_text(image_bytes: bytes, mime: str = "image/png") -> str:
    """Transcripción/extracción de datos de una imagen con visión de OpenAI.
    Corre en el subproceso de ingesta (cliente síncrono)."""
    import base64

    from openai import OpenAI

    from ...core.config import settings

    client = OpenAI(api_key=settings.openai_api_key, timeout=180)
    b64 = base64.b64encode(image_bytes).decode()
    try:
        resp = client.chat.completions.create(
            model=settings.agent_model,
            messages=[
                {
                    "role": "user",
                    "content": [
                        {"type": "text", "text": _OCR_PROMPT},
                        {"type": "image_url", "image_url": {"url": f"data:{mime};base64,{b64}"}},
                    ],
                }
            ],
        )
        text = (resp.choices[0].message.content or "").strip()
        return "" if text.upper() == "VACIO" else text
    except Exception:
        return ""


def _vision_ocr(path: str, max_pages: int = 30) -> list[Section]:
    """OCR de PDF escaneado: cada página se renderiza a PNG y se transcribe."""
    import fitz  # PyMuPDF

    if not _vision_available():
        return []
    doc = fitz.open(path)
    pages: list[Section] = []
    for i, page in enumerate(doc, start=1):
        if i > max_pages:
            break
        pix = page.get_pixmap(dpi=150)
        text = _vision_image_text(pix.tobytes("png"))
        if text:
            pages.append({"text": text, "meta": {"page": i, "ocr": True}})
    doc.close()
    return pages


def extract_image(path: str, mime: str) -> list[Section]:
    """Imagen subida como fuente: se transcribe/analiza con visión."""
    if not _vision_available():
        raise ValueError(
            "Para indexar imágenes se necesita la IA de visión (OPENAI_API_KEY + OCR habilitado)."
        )
    data = Path(path).read_bytes()
    if len(data) > 18 * 1024 * 1024:
        raise ValueError("La imagen supera los 18 MB; reducila e intentá de nuevo.")
    text = _vision_image_text(data, mime or "image/png")
    return [{"text": text, "meta": {"ocr": True}}] if text else []


def _docx_altchunks(path: str) -> list[Section]:
    """Contenido en altChunk: documentos generados por conversores guardan el
    cuerpo real como HTML/MHT incrustado (word/afchunk.mht) y el document.xml
    queda vacío. Se parsea el MIME y se extrae el HTML."""
    import zipfile

    sections: list[Section] = []
    try:
        with zipfile.ZipFile(path) as z:
            chunks = [
                n for n in z.namelist()
                if n.startswith("word/") and n.lower().endswith((".mht", ".mhtml", ".html", ".htm"))
            ]
            for name in chunks[:5]:
                raw = z.read(name)
                html_text = ""
                if name.lower().endswith((".mht", ".mhtml")):
                    import email

                    msg = email.message_from_bytes(raw)
                    for part in msg.walk():
                        if part.get_content_type() == "text/html":
                            payload = part.get_payload(decode=True) or b""
                            charset = part.get_content_charset() or "utf-8"
                            html_text = payload.decode(charset, errors="replace")
                            break
                    if not html_text and raw.lstrip()[:1] == b"<":
                        html_text = raw.decode("utf-8", errors="replace")
                else:
                    html_text = raw.decode("utf-8", errors="replace")

                if html_text:
                    from bs4 import BeautifulSoup

                    soup = BeautifulSoup(html_text, "html.parser")
                    for tag in soup(["script", "style"]):
                        tag.decompose()
                    text = "\n".join(
                        t.strip() for t in soup.get_text("\n").splitlines() if t.strip()
                    )
                    if text:
                        sections.append({"text": text, "meta": {"section": "contenido"}})
    except Exception:
        pass
    return sections


def _docx_textboxes(path: str) -> list[str]:
    """Texto dentro de cuadros de texto/formas (python-docx no los lee)."""
    import zipfile
    from xml.etree import ElementTree

    W = "{http://schemas.openxmlformats.org/wordprocessingml/2006/main}"
    texts: list[str] = []
    try:
        with zipfile.ZipFile(path) as z:
            xml = z.read("word/document.xml")
        root = ElementTree.fromstring(xml)
        for txbx in root.iter(f"{W}txbxContent"):
            fragment = "".join(node.text or "" for node in txbx.iter(f"{W}t")).strip()
            if fragment:
                texts.append(fragment)
    except Exception:
        pass
    return texts


def _docx_images_ocr(path: str, max_images: int = 15) -> list[Section]:
    """OCR de las imágenes incrustadas en el Word (documentos armados con capturas)."""
    import zipfile

    if not _vision_available():
        return []
    sections: list[Section] = []
    mime_by_ext = {".png": "image/png", ".jpg": "image/jpeg", ".jpeg": "image/jpeg",
                   ".webp": "image/webp", ".gif": "image/gif"}
    try:
        with zipfile.ZipFile(path) as z:
            media = [
                n for n in z.namelist()
                if n.startswith("word/media/") and Path(n).suffix.lower() in mime_by_ext
            ][:max_images]
            for idx, name in enumerate(media, start=1):
                data = z.read(name)
                if len(data) < 8_000:  # íconos y decoraciones
                    continue
                text = _vision_image_text(data, mime_by_ext[Path(name).suffix.lower()])
                if text:
                    sections.append({"text": text, "meta": {"section": f"imagen {idx}", "ocr": True}})
    except Exception:
        pass
    return sections


def extract_docx(path: str) -> list[Section]:
    import docx

    document = docx.Document(path)
    sections: list[Section] = []
    current_heading = ""
    buffer: list[str] = []

    def flush() -> None:
        if buffer:
            sections.append({
                "text": "\n".join(buffer).strip(),
                "meta": {"section": current_heading or None},
            })
            buffer.clear()

    for para in document.paragraphs:
        text = para.text.strip()
        if not text:
            continue
        if para.style.name.lower().startswith("heading"):
            flush()
            current_heading = text
            buffer.append(f"## {text}")
        else:
            buffer.append(text)
    flush()

    for t_idx, table in enumerate(document.tables, start=1):
        rows = []
        for row in table.rows:
            rows.append(" | ".join(c.text.strip() for c in row.cells))
        if rows:
            sections.append({
                "text": "\n".join(rows),
                "meta": {"section": f"tabla {t_idx}"},
            })
    sections = [s for s in sections if s["text"]]

    total_chars = sum(len(s["text"]) for s in sections)
    if total_chars < 200:
        # Contenido en altChunk (HTML/MHT incrustado por conversores)
        sections.extend(_docx_altchunks(path))
        total_chars = sum(len(s["text"]) for s in sections)
    if total_chars < 200:
        # Words armados con cuadros de texto (python-docx no los lee)
        for fragment in _docx_textboxes(path):
            sections.append({"text": fragment, "meta": {"section": "cuadro de texto"}})
        total_chars = sum(len(s["text"]) for s in sections)
    if total_chars < 200:
        # Words armados con imágenes/capturas: OCR con visión
        sections.extend(_docx_images_ocr(path))
    return [s for s in sections if s["text"]]


def extract_xlsx(path: str, max_rows: int = 200000) -> list[Section]:
    """Cada hoja se convierte a tabla markdown, troceada en bloques de filas
    que conservan el encabezado (citables por hoja y rango)."""
    from openpyxl import load_workbook

    wb = load_workbook(path, read_only=True, data_only=True)
    sections: list[Section] = []
    BLOCK = 40  # filas por bloque

    for sheet in wb.worksheets:
        rows_iter = sheet.iter_rows(values_only=True)
        try:
            header = next(rows_iter)
        except StopIteration:
            continue
        header_cells = [str(c) if c is not None else "" for c in header]
        header_md = "| " + " | ".join(header_cells) + " |"
        sep_md = "|" + "---|" * len(header_cells)

        block: list[str] = []
        start_row = 2
        count = 0
        for r_idx, row in enumerate(rows_iter, start=2):
            if count >= max_rows:
                break
            cells = [str(c) if c is not None else "" for c in row]
            if not any(c.strip() for c in cells):
                continue
            block.append("| " + " | ".join(cells) + " |")
            count += 1
            if len(block) >= BLOCK:
                sections.append({
                    "text": "\n".join([header_md, sep_md, *block]),
                    "meta": {"sheet": sheet.title, "rows": f"{start_row}-{r_idx}"},
                })
                block = []
                start_row = r_idx + 1
        if block:
            sections.append({
                "text": "\n".join([header_md, sep_md, *block]),
                "meta": {"sheet": sheet.title, "rows": f"{start_row}+"},
            })
    wb.close()
    return sections


def _assert_public_http_url(url: str) -> None:
    """Anti-SSRF: solo http(s) hacia hosts PÚBLICOS.

    Rechaza loopback, redes privadas, link-local (incluida la metadata de
    nube 169.254.169.254) y reservadas, resolviendo el hostname."""
    import ipaddress
    import socket
    from urllib.parse import urlparse

    parsed = urlparse(url)
    if parsed.scheme not in ("http", "https") or not parsed.hostname:
        raise ValueError("URL inválida: solo se admiten http/https")
    try:
        infos = socket.getaddrinfo(parsed.hostname, None)
    except OSError as exc:
        raise ValueError(f"No se pudo resolver el host: {parsed.hostname}") from exc
    for info in infos:
        ip = ipaddress.ip_address(info[4][0])
        if (ip.is_private or ip.is_loopback or ip.is_link_local
                or ip.is_reserved or ip.is_multicast or ip.is_unspecified):
            raise ValueError("La URL apunta a una red interna: no permitida")


def extract_link(url: str) -> tuple[list[Section], str]:
    """Descarga la página y extrae el contenido principal con trafilatura.

    Los redirects se siguen manualmente re-validando cada salto (un redirect
    hacia una IP interna no puede saltarse el control anti-SSRF)."""
    import httpx

    current = url
    resp = None
    for _ in range(4):  # hasta 3 redirects
        _assert_public_http_url(current)
        resp = httpx.get(
            current,
            follow_redirects=False,
            timeout=45,
            headers={"User-Agent": "Mozilla/5.0 (compatible; VEXConsulting/1.0)"},
        )
        if resp.status_code in (301, 302, 303, 307, 308) and resp.headers.get("location"):
            current = str(resp.next_request.url) if resp.next_request else resp.headers["location"]
            continue
        break
    resp.raise_for_status()
    html = resp.text

    text = ""
    title = url
    try:
        import trafilatura

        extracted = trafilatura.extract(html, include_tables=True, include_links=False)
        if extracted:
            text = extracted
        meta = trafilatura.extract_metadata(html)
        if meta and meta.title:
            title = meta.title
    except Exception:
        pass

    if not text:
        from bs4 import BeautifulSoup

        soup = BeautifulSoup(html, "html.parser")
        for tag in soup(["script", "style", "nav", "footer", "header"]):
            tag.decompose()
        text = "\n".join(t.strip() for t in soup.get_text("\n").splitlines() if t.strip())
        if soup.title and soup.title.string:
            title = soup.title.string.strip()

    return ([{"text": text, "meta": {"section": "contenido"}}] if text else []), title


def extract_text_file(path: str) -> list[Section]:
    raw = Path(path).read_text(encoding="utf-8", errors="replace")
    return [{"text": raw, "meta": {}}] if raw.strip() else []


def extract_source(kind: str, stored_path: str | None, url: str | None,
                   mime_type: str | None, max_rows: int) -> dict:
    """Dispatch principal. Devuelve {sections, page_count, title_hint}."""
    if kind == "link" and url:
        sections, title = extract_link(url)
        return {"sections": sections, "page_count": None, "title_hint": title}

    if not stored_path:
        return {"sections": [], "page_count": None, "title_hint": None}

    mt = (mime_type or "").lower()
    lower = stored_path.lower()
    if "pdf" in mt or lower.endswith(".pdf"):
        pages, n = extract_pdf(stored_path)
        return {"sections": pages, "page_count": n, "title_hint": None}
    if lower.endswith(".docx") or "wordprocessingml" in mt:
        return {"sections": extract_docx(stored_path), "page_count": None, "title_hint": None}
    if lower.endswith((".xlsx", ".xlsm")) or "spreadsheetml" in mt:
        return {"sections": extract_xlsx(stored_path, max_rows), "page_count": None, "title_hint": None}
    if lower.endswith((".txt", ".md", ".csv")) or mt.startswith("text/"):
        return {"sections": extract_text_file(stored_path), "page_count": None, "title_hint": None}
    if mt.startswith("image/") or lower.endswith((".png", ".jpg", ".jpeg", ".webp", ".gif")):
        return {"sections": extract_image(stored_path, mt), "page_count": None, "title_hint": None}

    raise ValueError(
        "Formato no soportado. Aceptamos PDF, Word (.docx), Excel (.xlsx), "
        "texto (.txt/.md/.csv), imágenes (PNG/JPG) y links."
    )
